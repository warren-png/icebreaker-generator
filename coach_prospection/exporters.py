"""
Génération de livrables Word (.docx) et PDF à partir des sorties Markdown
de Claude. Rendu type "document Entourage" : titres Playfair, accent or,
mise en page exécutive.

Usage :
    from coach_prospection.exporters import markdown_to_docx, markdown_to_pdf

    docx_bytes = markdown_to_docx(claude_output, title="Débrief — appel Marie Dupont")
    pdf_bytes = markdown_to_pdf(claude_output, title="Débrief — appel Marie Dupont")
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

GOLD = "#C9A227"          # or sobre, exécutif
GOLD_RGB = (201, 162, 39)
INK = "#0A0A0A"           # quasi-noir
INK_RGB = (10, 10, 10)
MUTED = "#5A5A5A"
PAPER = "#FAFAF8"


# ===========================================================================
# Parsing du markdown produit par Claude
# ===========================================================================

def _strip_inline_markdown(text: str) -> str:
    """Enlève les marqueurs gras/italique pour rendu plain text avec runs séparés."""
    # cette fonction n'est PAS utilisée pour le rendu : on parse les runs séparément.
    return text


def _parse_runs(text: str) -> list[tuple[str, dict]]:
    """Découpe une ligne markdown en runs avec leurs styles.

    Reconnaît : **gras**, *italique*, ***gras italique***, `code`, > blockquote.
    Retourne une liste de tuples (texte, attrs) où attrs = {bold, italic, code}.
    """
    runs: list[tuple[str, dict]] = []
    # Pattern combiné pour ***/**/*/`...`
    pattern = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos : m.start()], {}))
        tok = m.group(0)
        if tok.startswith("***") and tok.endswith("***"):
            runs.append((tok[3:-3], {"bold": True, "italic": True}))
        elif tok.startswith("**") and tok.endswith("**"):
            runs.append((tok[2:-2], {"bold": True}))
        elif tok.startswith("`") and tok.endswith("`"):
            runs.append((tok[1:-1], {"code": True}))
        elif tok.startswith("*") and tok.endswith("*"):
            runs.append((tok[1:-1], {"italic": True}))
        else:
            runs.append((tok, {}))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], {}))
    return runs


def _classify_block(line: str) -> tuple[str, str]:
    """Détecte le type d'un bloc markdown (heading, bullet, blockquote, paragraph)."""
    s = line.rstrip()
    if not s.strip():
        return ("blank", "")
    if s.startswith("### "):
        return ("h3", s[4:].strip())
    if s.startswith("## "):
        return ("h2", s[3:].strip())
    if s.startswith("# "):
        return ("h1", s[2:].strip())
    if s.startswith("> "):
        return ("quote", s[2:].strip())
    if re.match(r"^\s*[-*+]\s+", s):
        return ("bullet", re.sub(r"^\s*[-*+]\s+", "", s))
    if re.match(r"^\s*\d+\.\s+", s):
        return ("numbered", re.sub(r"^\s*\d+\.\s+", "", s))
    if s.strip() in ("---", "***"):
        return ("hr", "")
    return ("p", s.strip())


# ===========================================================================
# Génération Word (.docx)
# ===========================================================================

def _set_run_style(run, *, bold=False, italic=False, color=None, size=None, font=None):
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font)
        rfonts.set(qn("w:hAnsi"), font)
        rpr.append(rfonts)


def _add_runs_to_paragraph(para, runs: list[tuple[str, dict]], base_font="Calibri", base_size=11, base_color=INK_RGB):
    for text, attrs in runs:
        run = para.add_run(text)
        _set_run_style(
            run,
            bold=attrs.get("bold", False),
            italic=attrs.get("italic", False),
            color=base_color,
            size=base_size,
            font="Consolas" if attrs.get("code") else base_font,
        )


def _add_horizontal_rule(doc, color="C9A227"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def markdown_to_docx(markdown_text: str, title: str, subtitle: str | None = None) -> bytes:
    """Convertit le markdown de Claude en .docx Entourage."""
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Bandeau de tête : titre + date
    head_para = doc.add_paragraph()
    head_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = head_para.add_run("ENTOURAGE RECRUTEMENT")
    _set_run_style(run, bold=True, color=GOLD_RGB, size=9, font="Calibri")
    run.font.all_caps = True

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    _set_run_style(title_run, bold=True, color=INK_RGB, size=22, font="Cambria")

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(subtitle)
        _set_run_style(sub_run, italic=True, color=(90, 90, 90), size=11, font="Calibri")

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime("%d %B %Y"))
    _set_run_style(date_run, color=(120, 120, 120), size=9, font="Calibri")

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # Corps
    in_blockquote_block = False
    for raw_line in markdown_text.splitlines():
        kind, payload = _classify_block(raw_line)

        if kind == "blank":
            in_blockquote_block = False
            continue
        if kind == "hr":
            _add_horizontal_rule(doc, color="DDDDDD")
            continue
        if kind == "h1":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            _set_run_style(r, bold=True, color=INK_RGB, size=18, font="Cambria")
            continue
        if kind == "h2":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            _set_run_style(r, bold=True, color=INK_RGB, size=14, font="Cambria")
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            continue
        if kind == "h3":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            _set_run_style(r, bold=True, color=GOLD_RGB, size=11, font="Calibri")
            r.font.all_caps = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            continue
        if kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            # Bordure gauche
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "16")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "C9A227")
            pBdr.append(left)
            pPr.append(pBdr)
            _add_runs_to_paragraph(p, _parse_runs(payload), base_size=11, base_color=(60, 60, 60))
            in_blockquote_block = True
            continue
        if kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_to_paragraph(p, _parse_runs(payload))
            continue
        if kind == "numbered":
            p = doc.add_paragraph(style="List Number")
            _add_runs_to_paragraph(p, _parse_runs(payload))
            continue
        # paragraph
        p = doc.add_paragraph()
        _add_runs_to_paragraph(p, _parse_runs(payload))
        p.paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# Génération PDF (via reportlab)
# ===========================================================================

def _pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "EntTitle",
            parent=base["Heading1"],
            fontName="Times-Bold",
            fontSize=20,
            textColor=HexColor(INK),
            spaceAfter=4,
            leading=24,
        ),
        "subtitle": ParagraphStyle(
            "EntSubtitle",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=10,
            textColor=HexColor(MUTED),
            spaceAfter=2,
        ),
        "date": ParagraphStyle(
            "EntDate",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8,
            textColor=HexColor(MUTED),
            spaceAfter=12,
        ),
        "h1": ParagraphStyle(
            "EntH1",
            parent=base["Heading1"],
            fontName="Times-Bold",
            fontSize=16,
            textColor=HexColor(INK),
            spaceBefore=14,
            spaceAfter=6,
            leading=20,
        ),
        "h2": ParagraphStyle(
            "EntH2",
            parent=base["Heading2"],
            fontName="Times-Bold",
            fontSize=13,
            textColor=HexColor(INK),
            spaceBefore=10,
            spaceAfter=4,
            leading=16,
        ),
        "h3": ParagraphStyle(
            "EntH3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=9,
            textColor=HexColor(GOLD),
            spaceBefore=8,
            spaceAfter=3,
            leading=12,
        ),
        "body": ParagraphStyle(
            "EntBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            textColor=HexColor(INK),
            leading=14,
            spaceAfter=4,
            alignment=4,  # justified
        ),
        "quote": ParagraphStyle(
            "EntQuote",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=10,
            textColor=HexColor(MUTED),
            leftIndent=14,
            borderColor=HexColor(GOLD),
            borderWidth=0,
            borderPadding=0,
            leading=14,
            spaceAfter=4,
        ),
        "bullet": ParagraphStyle(
            "EntBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            textColor=HexColor(INK),
            leftIndent=14,
            bulletIndent=2,
            leading=14,
            spaceAfter=2,
        ),
    }


def _md_runs_to_pdf_markup(text: str) -> str:
    """Convertit le markdown inline en balises reportlab (<b>, <i>, <font>)."""
    # Préserve les espaces ; échappe les caractères HTML hostiles minimaux
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # ***bold italic***
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", text)
    # **bold**
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # *italic*
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    # `code`
    text = re.sub(r"`(.+?)`", r'<font face="Courier">\1</font>', text)
    return text


def markdown_to_pdf(markdown_text: str, title: str, subtitle: str | None = None) -> bytes:
    """Convertit le markdown en PDF style document Entourage."""
    buf = io.BytesIO()
    styles = _pdf_styles()

    def _on_page(canvas, doc):
        canvas.saveState()
        # Bandeau or en haut
        canvas.setFillColor(HexColor(GOLD))
        canvas.rect(0, A4[1] - 4 * mm, A4[0], 4 * mm, fill=1, stroke=0)
        # Petit footer
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(HexColor(MUTED))
        canvas.drawCentredString(
            A4[0] / 2,
            10 * mm,
            f"Entourage Recrutement — Coach Prospection — page {doc.page}",
        )
        canvas.restoreState()

    pdf = BaseDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
        author="Entourage Recrutement",
    )
    frame = Frame(
        pdf.leftMargin,
        pdf.bottomMargin,
        pdf.width,
        pdf.height,
        showBoundary=0,
    )
    pdf.addPageTemplates([PageTemplate(id="entourage", frames=frame, onPage=_on_page)])

    story = []
    # En-tête doc
    story.append(Paragraph("ENTOURAGE RECRUTEMENT", styles["h3"]))
    story.append(Paragraph(title, styles["title"]))
    if subtitle:
        story.append(Paragraph(subtitle, styles["subtitle"]))
    story.append(Paragraph(datetime.now().strftime("%d %B %Y"), styles["date"]))

    # Séparateur or
    story.append(
        Table(
            [[""]],
            colWidths=[pdf.width],
            rowHeights=[1.5],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), HexColor(GOLD)),
                    ("BOX", (0, 0), (-1, -1), 0, HexColor(GOLD)),
                ]
            ),
        )
    )
    story.append(Spacer(1, 6))

    bullet_buffer: list[str] = []

    def _flush_bullets():
        for b in bullet_buffer:
            story.append(Paragraph(f"•&nbsp;&nbsp;{_md_runs_to_pdf_markup(b)}", styles["bullet"]))
        bullet_buffer.clear()

    for raw_line in markdown_text.splitlines():
        kind, payload = _classify_block(raw_line)

        if kind == "blank":
            _flush_bullets()
            continue
        if kind in {"bullet", "numbered"}:
            bullet_buffer.append(payload)
            continue
        else:
            _flush_bullets()

        if kind == "hr":
            story.append(Spacer(1, 4))
            story.append(
                Table(
                    [[""]],
                    colWidths=[pdf.width],
                    rowHeights=[0.6],
                    style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), HexColor("#DDDDDD"))]),
                )
            )
            story.append(Spacer(1, 4))
        elif kind == "h1":
            story.append(Paragraph(_md_runs_to_pdf_markup(payload), styles["h1"]))
        elif kind == "h2":
            story.append(Paragraph(_md_runs_to_pdf_markup(payload), styles["h2"]))
        elif kind == "h3":
            story.append(Paragraph(_md_runs_to_pdf_markup(payload), styles["h3"]))
        elif kind == "quote":
            story.append(Paragraph(_md_runs_to_pdf_markup(payload), styles["quote"]))
        else:  # paragraph
            story.append(Paragraph(_md_runs_to_pdf_markup(payload), styles["body"]))

    _flush_bullets()
    pdf.build(story)
    return buf.getvalue()
